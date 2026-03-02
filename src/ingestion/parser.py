"""Document parsers for PDF, DOCX, Markdown, and TXT file formats."""

import hashlib
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """Metadata extracted from a parsed document.

    Attributes:
        source_file: Original file path.
        title: Document title (derived from filename).
        created_date: Timestamp when parsing occurred.
        page_count: Number of pages (1 for non-paginated formats).
        file_hash: MD5 hash of the file contents for deduplication.
        file_type: File extension without dot.
    """

    source_file: str
    title: str
    created_date: str
    page_count: int
    file_hash: str
    file_type: str


@dataclass
class ParsedDocument:
    """Result of parsing a document file.

    Attributes:
        content: Full text content of the document.
        metadata: Extracted metadata.
        pages: List of per-page text content.
    """

    content: str
    metadata: DocumentMetadata
    pages: list[str] = field(default_factory=list)


def compute_file_hash(file_path: Path) -> str:
    """Compute MD5 hash of a file for duplicate detection.

    Args:
        file_path: Path to the file.

    Returns:
        Hex digest of the file's MD5 hash.
    """
    hasher = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


class BaseParser(ABC):
    """Abstract base class for document parsers."""

    @abstractmethod
    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a document file into structured content.

        Args:
            file_path: Path to the document file.

        Returns:
            ParsedDocument with extracted text and metadata.
        """

    def _build_metadata(self, file_path: Path, page_count: int) -> DocumentMetadata:
        """Build metadata for a parsed document.

        Args:
            file_path: Path to the source file.
            page_count: Number of pages in the document.

        Returns:
            DocumentMetadata instance.
        """
        return DocumentMetadata(
            source_file=str(file_path),
            title=file_path.stem,
            created_date=datetime.now().isoformat(),
            page_count=page_count,
            file_hash=compute_file_hash(file_path),
            file_type=file_path.suffix.lstrip(".").lower(),
        )


class PDFParser(BaseParser):
    """Parser for PDF documents using PyPDF2."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a PDF file into text content.

        Args:
            file_path: Path to the PDF file.

        Returns:
            ParsedDocument with per-page text extraction.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file is not a valid PDF.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        try:
            reader = PdfReader(str(file_path))
        except Exception as e:
            raise ValueError(f"Failed to read PDF: {file_path}") from e

        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)

        content = "\n\n".join(pages)
        metadata = self._build_metadata(file_path, len(pages))
        logger.info("Parsed PDF: %s (%d pages)", file_path.name, len(pages))

        return ParsedDocument(content=content, metadata=metadata, pages=pages)


class DOCXParser(BaseParser):
    """Parser for DOCX documents using python-docx."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a DOCX file into text content.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            ParsedDocument with paragraph-based content.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file is not a valid DOCX.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            doc = DocxDocument(str(file_path))
        except Exception as e:
            raise ValueError(f"Failed to read DOCX: {file_path}") from e

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        metadata = self._build_metadata(file_path, 1)
        logger.info(
            "Parsed DOCX: %s (%d paragraphs)",
            file_path.name,
            len(paragraphs),
        )

        return ParsedDocument(content=content, metadata=metadata, pages=[content])


class MarkdownParser(BaseParser):
    """Parser for Markdown files with HTML stripping."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a Markdown file into plain text content.

        Args:
            file_path: Path to the Markdown file.

        Returns:
            ParsedDocument with markdown content preserved as-is.

        Raises:
            FileNotFoundError: If the file does not exist.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Markdown file not found: {file_path}")

        raw = file_path.read_text(encoding="utf-8")
        content = raw.strip()
        metadata = self._build_metadata(file_path, 1)
        logger.info("Parsed Markdown: %s", file_path.name)

        return ParsedDocument(content=content, metadata=metadata, pages=[content])


class TextParser(BaseParser):
    """Parser for plain text files."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a plain text file.

        Args:
            file_path: Path to the text file.

        Returns:
            ParsedDocument with raw text content.

        Raises:
            FileNotFoundError: If the file does not exist.
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Text file not found: {file_path}")

        content = file_path.read_text(encoding="utf-8").strip()
        metadata = self._build_metadata(file_path, 1)
        logger.info("Parsed TXT: %s", file_path.name)

        return ParsedDocument(content=content, metadata=metadata, pages=[content])


_PARSER_MAP: dict[str, type[BaseParser]] = {
    "pdf": PDFParser,
    "docx": DOCXParser,
    "md": MarkdownParser,
    "txt": TextParser,
}


class ParserFactory:
    """Factory for selecting the appropriate parser based on file type."""

    @staticmethod
    def get_parser(file_path: str | Path) -> BaseParser:
        """Return the correct parser for a given file.

        Args:
            file_path: Path to the document file.

        Returns:
            An instance of the appropriate parser.

        Raises:
            ValueError: If the file format is not supported.
        """
        ext = Path(file_path).suffix.lstrip(".").lower()
        parser_cls = _PARSER_MAP.get(ext)
        if parser_cls is None:
            supported = ", ".join(sorted(_PARSER_MAP.keys()))
            raise ValueError(f"Unsupported format: .{ext}. Supported: {supported}")
        return parser_cls()

    @staticmethod
    def supported_formats() -> list[str]:
        """Return list of supported file extensions.

        Returns:
            Sorted list of supported format strings.
        """
        return sorted(_PARSER_MAP.keys())
