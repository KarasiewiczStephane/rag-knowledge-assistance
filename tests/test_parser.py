"""Tests for document parsers (PDF, DOCX, Markdown, TXT)."""

from pathlib import Path

import pytest
from docx import Document as DocxDocument
from PyPDF2 import PdfWriter

from src.ingestion.parser import (
    DOCXParser,
    MarkdownParser,
    ParserFactory,
    PDFParser,
    TextParser,
    compute_file_hash,
)


@pytest.fixture()
def sample_txt(tmp_path: Path) -> Path:
    """Create a sample text file."""
    f = tmp_path / "sample.txt"
    f.write_text("Hello, this is a test document.\nSecond line.")
    return f


@pytest.fixture()
def sample_md(tmp_path: Path) -> Path:
    """Create a sample markdown file."""
    f = tmp_path / "sample.md"
    f.write_text("# Title\n\nSome **bold** text.\n\n## Section\n\nMore content.")
    return f


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    """Create a minimal PDF file."""
    f = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with open(f, "wb") as out:
        writer.write(out)
    return f


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Create a sample DOCX file."""
    f = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph of the document.")
    doc.add_paragraph("Second paragraph with more content.")
    doc.save(str(f))
    return f


def test_compute_file_hash(sample_txt: Path) -> None:
    """File hash is a 32-char hex string."""
    h = compute_file_hash(sample_txt)
    assert isinstance(h, str)
    assert len(h) == 32


def test_compute_file_hash_deterministic(sample_txt: Path) -> None:
    """Same file produces the same hash."""
    h1 = compute_file_hash(sample_txt)
    h2 = compute_file_hash(sample_txt)
    assert h1 == h2


def test_txt_parser(sample_txt: Path) -> None:
    """TextParser extracts content and metadata."""
    parser = TextParser()
    result = parser.parse(sample_txt)
    assert "Hello" in result.content
    assert result.metadata.file_type == "txt"
    assert result.metadata.page_count == 1
    assert len(result.pages) == 1


def test_txt_parser_missing_file() -> None:
    """TextParser raises FileNotFoundError for missing files."""
    parser = TextParser()
    with pytest.raises(FileNotFoundError):
        parser.parse(Path("/nonexistent/file.txt"))


def test_md_parser(sample_md: Path) -> None:
    """MarkdownParser extracts content and metadata."""
    parser = MarkdownParser()
    result = parser.parse(sample_md)
    assert "Title" in result.content
    assert "bold" in result.content
    assert result.metadata.file_type == "md"
    assert result.metadata.page_count == 1


def test_md_parser_missing_file() -> None:
    """MarkdownParser raises FileNotFoundError for missing files."""
    parser = MarkdownParser()
    with pytest.raises(FileNotFoundError):
        parser.parse(Path("/nonexistent/file.md"))


def test_pdf_parser(sample_pdf: Path) -> None:
    """PDFParser extracts pages from a PDF."""
    parser = PDFParser()
    result = parser.parse(sample_pdf)
    assert result.metadata.file_type == "pdf"
    assert result.metadata.page_count == 1
    assert len(result.pages) == 1


def test_pdf_parser_missing_file() -> None:
    """PDFParser raises FileNotFoundError for missing files."""
    parser = PDFParser()
    with pytest.raises(FileNotFoundError):
        parser.parse(Path("/nonexistent/file.pdf"))


def test_pdf_parser_invalid_file(tmp_path: Path) -> None:
    """PDFParser raises ValueError for an invalid PDF."""
    bad = tmp_path / "bad.pdf"
    bad.write_text("not a pdf")
    parser = PDFParser()
    with pytest.raises(ValueError, match="Failed to read PDF"):
        parser.parse(bad)


def test_docx_parser(sample_docx: Path) -> None:
    """DOCXParser extracts paragraphs from a DOCX."""
    parser = DOCXParser()
    result = parser.parse(sample_docx)
    assert "First paragraph" in result.content
    assert "Second paragraph" in result.content
    assert result.metadata.file_type == "docx"


def test_docx_parser_missing_file() -> None:
    """DOCXParser raises FileNotFoundError for missing files."""
    parser = DOCXParser()
    with pytest.raises(FileNotFoundError):
        parser.parse(Path("/nonexistent/file.docx"))


def test_docx_parser_invalid_file(tmp_path: Path) -> None:
    """DOCXParser raises ValueError for an invalid DOCX."""
    bad = tmp_path / "bad.docx"
    bad.write_text("not a docx")
    parser = DOCXParser()
    with pytest.raises(ValueError, match="Failed to read DOCX"):
        parser.parse(bad)


def test_parser_factory_txt(sample_txt: Path) -> None:
    """ParserFactory returns TextParser for .txt files."""
    parser = ParserFactory.get_parser(sample_txt)
    assert isinstance(parser, TextParser)


def test_parser_factory_md(sample_md: Path) -> None:
    """ParserFactory returns MarkdownParser for .md files."""
    parser = ParserFactory.get_parser(sample_md)
    assert isinstance(parser, MarkdownParser)


def test_parser_factory_pdf(sample_pdf: Path) -> None:
    """ParserFactory returns PDFParser for .pdf files."""
    parser = ParserFactory.get_parser(sample_pdf)
    assert isinstance(parser, PDFParser)


def test_parser_factory_docx(sample_docx: Path) -> None:
    """ParserFactory returns DOCXParser for .docx files."""
    parser = ParserFactory.get_parser(sample_docx)
    assert isinstance(parser, DOCXParser)


def test_parser_factory_unsupported() -> None:
    """ParserFactory raises ValueError for unsupported formats."""
    with pytest.raises(ValueError, match="Unsupported format"):
        ParserFactory.get_parser("file.xyz")


def test_parser_factory_supported_formats() -> None:
    """ParserFactory.supported_formats returns all supported types."""
    formats = ParserFactory.supported_formats()
    assert "pdf" in formats
    assert "docx" in formats
    assert "md" in formats
    assert "txt" in formats


def test_metadata_has_hash(sample_txt: Path) -> None:
    """Parsed document metadata includes a file hash."""
    parser = TextParser()
    result = parser.parse(sample_txt)
    assert len(result.metadata.file_hash) == 32


def test_metadata_has_title(sample_txt: Path) -> None:
    """Parsed document metadata uses filename stem as title."""
    parser = TextParser()
    result = parser.parse(sample_txt)
    assert result.metadata.title == "sample"
